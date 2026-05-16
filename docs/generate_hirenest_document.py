from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
DOCX_PATH = OUT_DIR / "HireNest_Project_Documentation.docx"
USE_CASE_IMG = OUT_DIR / "hirenest_use_case_diagram.png"
WORKFLOW_IMG = OUT_DIR / "hirenest_workflow_diagram.png"

PURPLE = "#7B2FF7"
PINK = "#F107A3"
DARK = "#111827"
MUTED = "#6B7280"
LIGHT_PURPLE = "#F3E8FF"
LIGHT_PINK = "#FCE7F3"
LIGHT_BLUE = "#E0F2FE"
LIGHT_GREEN = "#DCFCE7"
LIGHT_YELLOW = "#FEF3C7"
BORDER = "#D8B4FE"


def font(size=32, bold=False):
    candidates = [
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
    ]
    for path in candidates:
        if Path(path).exists():
            return ImageFont.truetype(path, size)
    return ImageFont.load_default()


def text_size(draw, text, fnt):
    box = draw.textbbox((0, 0), text, font=fnt)
    return box[2] - box[0], box[3] - box[1]


def center_text(draw, box, text, fnt, fill=DARK):
    x1, y1, x2, y2 = box
    w, h = text_size(draw, text, fnt)
    draw.text((x1 + (x2 - x1 - w) / 2, y1 + (y2 - y1 - h) / 2 - 2), text, font=fnt, fill=fill)


def rounded_box(draw, box, fill, outline="#C4B5FD", width=3, radius=24):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def arrow(draw, start, end, fill="#374151", width=4):
    draw.line([start, end], fill=fill, width=width)
    x1, y1 = start
    x2, y2 = end
    dx, dy = x2 - x1, y2 - y1
    if abs(dx) >= abs(dy):
        if dx > 0:
            pts = [(x2, y2), (x2 - 16, y2 - 10), (x2 - 16, y2 + 10)]
        else:
            pts = [(x2, y2), (x2 + 16, y2 - 10), (x2 + 16, y2 + 10)]
    else:
        if dy > 0:
            pts = [(x2, y2), (x2 - 10, y2 - 16), (x2 + 10, y2 - 16)]
        else:
            pts = [(x2, y2), (x2 - 10, y2 + 16), (x2 + 10, y2 + 16)]
    draw.polygon(pts, fill=fill)


def multiline_center(draw, box, lines, fnt, fill=DARK, gap=8):
    x1, y1, x2, y2 = box
    heights = [text_size(draw, line, fnt)[1] for line in lines]
    total = sum(heights) + gap * (len(lines) - 1)
    y = y1 + (y2 - y1 - total) / 2
    for line, h in zip(lines, heights):
        w, _ = text_size(draw, line, fnt)
        draw.text((x1 + (x2 - x1 - w) / 2, y), line, font=fnt, fill=fill)
        y += h + gap


def create_use_case_diagram():
    img = Image.new("RGB", (1900, 1250), "white")
    d = ImageDraw.Draw(img)
    title = font(48, True)
    heading = font(32, True)
    body = font(27)
    small = font(22)

    center_text(d, (0, 30, 1900, 95), "HireNest Job Portal Use Case Diagram", title, DARK)
    center_text(d, (0, 92, 1900, 135), "MERN Stack: MongoDB, Express.js, React.js, Node.js", font(30), MUTED)

    student_cases = ["Register / Login", "Manage Profile", "Search Jobs", "View Job Details", "Apply for Job", "Save Jobs", "Track Status", "Chat with HR"]
    hr_cases = ["Login", "Post Job", "Manage Jobs", "View Applications", "Review Candidate", "Update Status", "Schedule Interview", "Chat with Students"]
    admin_cases = ["Login", "Manage Users", "Approve / Reject HR", "Monitor Jobs", "Manage Applications", "Moderate Platform"]

    def actor_group(title_text, items, x, y, fill, accent):
        panel = (x, y, x + 520, y + 760)
        rounded_box(d, panel, fill, outline=accent, width=3, radius=28)
        actor = (x + 45, y + 35, x + 475, y + 105)
        rounded_box(d, actor, fill, outline=accent, width=3, radius=22)
        center_text(d, actor, title_text, heading)
        boxes = []
        last = actor
        for i, item in enumerate(items):
            box = (x + 75, y + 145 + i * 72, x + 445, y + 198 + i * 72)
            rounded_box(d, box, "#FFFFFF", outline="#CBD5E1", width=2, radius=28)
            center_text(d, box, item, body)
            arrow(d, ((last[0] + last[2]) // 2, last[3]), ((box[0] + box[2]) // 2, box[1]), accent, 2)
            last = box
            boxes.append(box)
        return panel, actor, boxes

    student_panel, _, _ = actor_group("Student", student_cases, 90, 230, LIGHT_BLUE, "#2563EB")
    hr_panel, _, _ = actor_group("HR / Recruiter", hr_cases, 690, 230, LIGHT_GREEN, "#16A34A")
    admin_panel, _, _ = actor_group("Admin", admin_cases, 1290, 230, LIGHT_YELLOW, "#B45309")

    service_box = (500, 1040, 1400, 1165)
    rounded_box(d, service_box, "#FAF5FF", outline="#A855F7", width=4, radius=30)
    multiline_center(
        d,
        service_box,
        ["HireNest Job Portal Core Services", "Authentication | Jobs | Applications | Uploads | Messaging | Notifications"],
        font(25, True),
        DARK,
        gap=10,
    )
    for panel in [student_panel, hr_panel, admin_panel]:
        arrow(d, ((panel[0] + panel[2]) // 2, panel[3]), (950, service_box[1]), "#64748B", 3)
    img.save(USE_CASE_IMG, quality=95)


def create_workflow_diagram():
    img = Image.new("RGB", (1800, 1300), "white")
    d = ImageDraw.Draw(img)
    title = font(48, True)
    heading = font(31, True)
    body = font(24)

    center_text(d, (0, 25, 1800, 90), "HireNest Overall Workflow", title, DARK)
    center_text(d, (0, 88, 1800, 130), "Student, HR / Recruiter, and Admin process flow", font(28), MUTED)

    columns = [
        ("Student Workflow", 120, LIGHT_BLUE, ["Register / Login", "Manage Profile", "Search Jobs", "View Job Details", "Apply for Job", "Upload Resume", "Track Status", "Chat with HR"]),
        ("HR / Recruiter Workflow", 650, LIGHT_GREEN, ["Login", "Post Job", "Manage Jobs", "View Applications", "Review Candidate", "Update Status", "Interview Update", "Chat with Student"]),
        ("Admin Workflow", 1180, LIGHT_YELLOW, ["Login", "Manage Users", "Approve / Reject HR", "Monitor Jobs", "Manage Applications", "Moderate Platform"]),
    ]

    start = (760, 160, 1040, 225)
    rounded_box(d, start, "#FAF5FF", outline="#A855F7", radius=35)
    center_text(d, start, "Start", heading)

    for title_text, x, fill, steps in columns:
        header = (x, 285, x + 430, 360)
        rounded_box(d, header, fill, outline="#94A3B8", radius=22)
        center_text(d, header, title_text, heading)
        arrow(d, (900, 225), (x + 215, 285), "#7C3AED", 3)

        last = header
        for i, step in enumerate(steps):
            y = 410 + i * 88
            box = (x + 45, y, x + 385, y + 58)
            rounded_box(d, box, "#FFFFFF", outline="#CBD5E1", radius=18)
            center_text(d, box, step, body)
            arrow(d, ((last[0] + last[2]) // 2, last[3]), (x + 215, box[1]), "#475569", 3)
            last = box

        end_y = 410 + len(steps) * 88
        end = (x + 110, end_y, x + 320, end_y + 58)
        rounded_box(d, end, fill, outline="#94A3B8", radius=35)
        center_text(d, end, "End", body)
        arrow(d, (x + 215, last[3]), (x + 215, end[1]), "#475569", 3)

    img.save(WORKFLOW_IMG, quality=95)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill.lstrip("#"))
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(10.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, rows, headers=None):
    table = doc.add_table(rows=1 if headers else 0, cols=len(headers or rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    if headers:
        for i, header in enumerate(headers):
            set_cell_text(table.rows[0].cells[i], header, True)
            set_cell_shading(table.rows[0].cells[i], LIGHT_PURPLE)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(17, 24, 39)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.08
    for name, size in [("Title", 24), ("Subtitle", 12), ("Heading 1", 16), ("Heading 2", 13)]:
        styles[name].font.name = "Arial"
        styles[name].font.size = Pt(size)
        styles[name].font.color.rgb = RGBColor(17, 24, 39)


def add_footer(section):
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("HireNest Job Portal Project Documentation")
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(107, 114, 128)


def create_doc():
    doc = Document()
    style_doc(doc)
    add_footer(doc.sections[0])

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("HireNest Job Portal")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor(123, 47, 247)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Project Documentation for Internship Certificate Submission")
    r.font.name = "Arial"
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(107, 114, 128)

    meta = [
        ("Project Name", "HireNest Job Portal"),
        ("Project Type", "Full Stack Web Application"),
        ("Technology Stack", "MongoDB, Express.js, React.js, Node.js"),
        ("User Roles", "Student, HR / Recruiter, Admin"),
        ("Prepared For", "Company / Internship Evaluation"),
    ]
    add_table(doc, meta, headers=["Field", "Details"])

    add_heading(doc, "1. Project Overview")
    doc.add_paragraph(
        "HireNest is a full-stack MERN job portal designed to connect students with recruiters and simplify the recruitment process. "
        "The platform provides role-based dashboards for students, HR/recruiters, and administrators. Students can browse jobs and apply for suitable opportunities, HR users can post jobs and manage candidates, and admins can monitor platform activity and approve HR accounts."
    )

    add_heading(doc, "2. Project Objectives")
    add_bullets(doc, [
        "Create a single platform where students can find and apply for jobs.",
        "Provide HR users with job posting, application review, and candidate management tools.",
        "Allow admins to manage users, approve HR accounts, and moderate platform activity.",
        "Implement secure login, role-based access, resume uploads, notifications, and real-time messaging.",
        "Build a responsive, modern user interface suitable for desktop and mobile devices.",
    ])

    add_heading(doc, "3. Technology Stack")
    tech_rows = [
        ("Frontend", "React.js, Vite, Bootstrap, CSS, Lucide React", "Responsive UI, pages, dashboards, forms, navigation"),
        ("Backend", "Node.js, Express.js", "REST APIs, business logic, middleware, server setup"),
        ("Database", "MongoDB, Mongoose", "Users, jobs, applications, messages, notifications"),
        ("Authentication", "JWT, bcryptjs", "Secure login, password hashing, protected routes"),
        ("File Upload", "Multer, Cloudinary utility", "Resume/profile file upload and storage URL handling"),
        ("Real-time", "Socket.io, socket.io-client", "Messaging and live notifications"),
        ("API Client", "Axios", "Frontend and backend API communication"),
        ("Deployment", "Vercel, Render", "Frontend and backend hosting"),
    ]
    add_table(doc, tech_rows, headers=["Layer", "Technology / Package", "Purpose"])

    add_heading(doc, "4. Main Modules")
    modules = [
        ("Authentication Module", "Registration, login, JWT token generation, password hashing, protected routes, role checks."),
        ("Student Module", "Profile management, job search, saved jobs, resume upload, applications, status tracking, messaging."),
        ("HR / Recruiter Module", "Company profile, job posting, job management, application review, interview updates, offer-letter handling."),
        ("Admin Module", "User management, HR approval/rejection, job monitoring, application monitoring, platform moderation."),
        ("Job Module", "Job creation, search, filters, expiry date, active/closed/expired status, skills and category handling."),
        ("Application Module", "Student applications, unique application records, HR review, status updates, interview and offer data."),
        ("Messaging and Notification Module", "Real-time chat and notifications using Socket.io."),
        ("File Upload Module", "Resume/profile upload support using Multer and Cloudinary configuration."),
    ]
    add_table(doc, modules, headers=["Module", "Description"])

    add_heading(doc, "5. User Roles and Responsibilities")
    roles = [
        ("Student", "Register/login, manage profile, upload resume, search jobs, save jobs, apply, track status, chat with HR."),
        ("HR / Recruiter", "Login, manage company profile, post jobs, manage jobs, view applications, shortlist candidates, send interview updates, chat with students."),
        ("Admin", "Login, manage users, approve/reject HR accounts, monitor jobs/applications, moderate platform activity."),
    ]
    add_table(doc, roles, headers=["Role", "Responsibilities"])

    add_heading(doc, "6. System Flow")
    add_numbered(doc, [
        "User registers or logs in to the HireNest platform.",
        "The backend verifies credentials and generates a JWT token.",
        "The frontend redirects the user according to their role: Student, HR, or Admin.",
        "Students search jobs and submit applications with resume details.",
        "HR users review applications and update candidate status.",
        "Admins manage users and approve or reject HR registrations.",
        "Messages and notifications help users communicate during the hiring process.",
    ])

    add_heading(doc, "7. Use Case Diagram")
    doc.add_paragraph("The use case diagram shows how Student, HR/Recruiter, and Admin interact with the HireNest system.")
    doc.add_picture(str(USE_CASE_IMG), width=Inches(6.9))

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_footer(doc.sections[-1])
    add_heading(doc, "8. Workflow Diagram")
    doc.add_paragraph("The workflow diagram explains the major process flow for each role in the platform.")
    doc.add_picture(str(WORKFLOW_IMG), width=Inches(6.9))

    add_heading(doc, "9. Database Collections")
    db_rows = [
        ("User", "Stores name, email, password hash, role, account status, student profile, company profile, saved jobs."),
        ("Job", "Stores title, description, company, location, industry, salary, category, skills, expiry date, status, postedBy."),
        ("Application", "Stores job, student, HR, resume URL, cover letter, status, interview details, offer-letter details."),
        ("Message", "Stores conversation messages between students and HR users."),
        ("Notification", "Stores user notifications related to applications, messages, and platform activity."),
    ]
    add_table(doc, db_rows, headers=["Collection", "Purpose"])

    add_heading(doc, "10. API Summary")
    api_rows = [
        ("Auth APIs", "Register, login, current user, protected authentication actions."),
        ("Job APIs", "Create jobs, fetch jobs, search/filter jobs, update jobs, manage job status."),
        ("Application APIs", "Apply for job, fetch applications, update status, interview and offer updates."),
        ("Admin APIs", "Manage users, approve/reject HR accounts, monitor jobs and applications."),
        ("Message APIs", "Send and retrieve messages between platform users."),
        ("Notification APIs", "Create, fetch, and mark notifications as read."),
        ("File APIs", "Upload files through Multer and store/access uploaded file URLs."),
    ]
    add_table(doc, api_rows, headers=["API Area", "Description"])

    add_heading(doc, "11. Security and Validation")
    add_bullets(doc, [
        "Passwords are hashed before storage using bcryptjs.",
        "JWT tokens are used to protect private routes and APIs.",
        "Role-based authorization restricts Student, HR, and Admin actions.",
        "Duplicate job applications are prevented using a unique job and student index.",
        "CORS configuration allows controlled frontend-backend communication.",
        "Environment variables are used for secrets such as database URI and JWT secret.",
    ])

    add_heading(doc, "12. Testing Summary")
    test_rows = [
        ("Authentication", "Checked registration, login, protected route access, and role-based redirection."),
        ("Student Flow", "Checked job browsing, application submission, saved jobs, profile/resume update, and status tracking."),
        ("HR Flow", "Checked job posting, application review, candidate status update, and interview flow."),
        ("Admin Flow", "Checked user management, HR approval/rejection, and monitoring features."),
        ("Responsive UI", "Checked mobile and desktop layouts for dashboards, navbar/sidebar, auth pages, and job pages."),
        ("Build", "Frontend production build can be verified using npm run build inside Client."),
    ]
    add_table(doc, test_rows, headers=["Area", "Testing Performed"])

    add_heading(doc, "13. Deployment")
    doc.add_paragraph(
        "The frontend is configured for Vercel deployment with SPA routing support. The backend is designed for Render deployment and requires MongoDB, JWT, Cloudinary, admin email, and admin password environment variables."
    )
    add_bullets(doc, [
        "Frontend: Vercel",
        "Backend: Render",
        "Database: MongoDB Atlas or local MongoDB",
        "Static assets and uploaded files: Cloudinary configuration",
    ])

    add_heading(doc, "14. Challenges Faced")
    add_bullets(doc, [
        "Managing separate access and dashboards for Student, HR, and Admin roles.",
        "Connecting frontend authentication state with backend JWT-protected APIs.",
        "Handling resume/profile uploads and storing file URLs securely.",
        "Keeping application status updates consistent between student and HR dashboards.",
        "Making the interface responsive for both desktop and mobile users.",
    ])

    add_heading(doc, "15. Future Enhancements")
    add_bullets(doc, [
        "Email notifications for application and interview updates.",
        "Advanced analytics dashboards for HR and admin users.",
        "Resume parsing and smart job recommendations.",
        "Calendar integration for interview scheduling.",
        "More detailed audit logs for admin moderation.",
    ])

    add_heading(doc, "16. Conclusion")
    doc.add_paragraph(
        "HireNest successfully demonstrates a complete MERN stack recruitment platform with real-world features such as secure authentication, role-based access, job posting, job applications, profile management, file uploads, messaging, notifications, and admin moderation. "
        "The project helped apply full-stack development concepts, database design, API development, frontend UI design, authentication, and deployment workflow in a practical application."
    )

    doc.add_paragraph()
    sign = add_table(doc, [
        ("Student Name", ""),
        ("Company Name", ""),
        ("Mentor / Guide", ""),
        ("Date", ""),
        ("Signature", ""),
    ], headers=["Submission Field", "Details"])
    for row in sign.rows:
        for cell in row.cells:
            cell.height = Inches(0.35)

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    create_use_case_diagram()
    create_workflow_diagram()
    create_doc()
    print(DOCX_PATH)
